from docx import Document
from docx.shared import inches

# Create a new document
x = Document()

# add some text
x.add_paragraph('Lifes is good!')

# Add image
# document.add_picture('image.jpg', width=inches(2))

# save the document using file handling
with open('name.docx', 'wb') as f:
          document.save(f)




